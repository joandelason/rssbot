from docx import Document
from docx.shared import Inches
import sys, os

import htmlReader

class DocxConverter:
    def __init__(self, title, content, sourceUrl, path):
        document = Document()
        document.add_heading(title, 0)
        document.add_paragraph( content )
        document.add_paragraph( '\n')  
        document.add_paragraph( 'Source: ' + sourceUrl)  
        if not os.path.exists(path):            
            os.makedirs(path)

        document.save(path+'/'+self.replaceForbiddenChars(title)+'.docx')

    #OS forbidden chars->   "*/:<>?\|%
    def replaceForbiddenChars(self, s):
        s = s.replace('"', '\'').replace('*', 'x').replace('*', 'x').replace('/', ' or ')
        s = s.replace(':', '-').replace('<', '-').replace('>', '-').replace('?', '')
        s = s.replace('\\', ' or ').replace('|', ' or ').replace('%', ' percent')
        
        return s



